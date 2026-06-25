import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_docx():
    docx_path = r"c:\Users\sapna\Downloads\verity-windows-build\ai-forensics-dashboard\blackhat_india_submission_proposal.docx"
    doc = docx.Document()
    
    def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
        run.font.name = name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

    p_header = doc.add_paragraph()
    set_font(p_header.add_run("BLACK HAT INDIA – SUBMISSION PROPOSAL"), size=14, bold=True, color=RGBColor(204, 0, 0))
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_title = doc.add_paragraph()
    set_font(p_title.add_run("Tinker Tailor LLM Spy: Reconstructing 'Deleted' Chats & Hijacking Sessions from Chromium LevelDB Caches with Verity"), size=18, bold=True, color=RGBColor(17, 17, 17))
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Proposal Info', level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    meta_info = [
        ("Title", "Tinker Tailor LLM Spy: Reconstructing 'Deleted' Chats & Hijacking Sessions from Chromium LevelDB Caches with Verity"),
        ("Status", "Submitted / For Board Review"),
        ("Session Type", "Briefings"),
        ("Speaker", "Dr. Sapna Vikram Mewundi (Lead Forensics Researcher, Verity Project)"),
        ("Tracks", "AI, ML, & Data Science; Threat Hunting & Incident Response"),
        ("Format", "40-Minute Briefings")
    ]
    
    for idx, (k, v) in enumerate(meta_info):
        row = table.rows[idx]
        row.cells[0].paragraphs[0].add_run(k).bold = True
        row.cells[1].paragraphs[0].add_run(v)
        
    doc.add_paragraph()
    
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph("It’s coming, and you aren’t ready—your first forensic investigation involving Large Language Model (LLM) data leakage. GenAI portals (like ChatGPT, Claude, and Gemini) have become standard corporate utilities. To speed up workflows, developers and executives routinely copy and paste sensitive proprietary source code, internal configurations, and credentials into these clients. But when users click 'Delete Chat' or clear histories, does that data actually get erased?")
    doc.add_paragraph("The truth is far more persistent. Chromium-based browsers (Google Chrome, Microsoft Edge) and native Electron desktop wrappers cache conversation telemetry inside client-side IndexedDB databases, which are backed by Google's LevelDB storage engine. Since LevelDB uses an append-only log-structured merge-tree (LSM tree), deleted records remain intact in write-ahead logs (.log) and uncompacted Sorted String Tables (.sst/.ldb) on local disks for days or weeks.")
    doc.add_paragraph("In this talk, we reveal the low-level serialization structures of client-side LLM databases and how to carve this 'deleted' telemetry. We reverse engineer the V8 deserialization format (ValueSerializer), specifically dissecting V8 string tags (OneByteString vs TwoByteString UTF-16LE anomalies) and nested object boundary parsing. We introduce Verity, an open-source, zero-dependency python framework that dynamically sweeps all user profiles, bypasses active Windows file locks, and reconstructs deleted histories. Finally, we establish a cryptographically secured chain of custody using in-browser HMAC-SHA256 verification. Attendees will leave with a ready-to-use incident response playbook and the open-source Verity tool to audit and protect their organizations.")

    doc.add_heading('Presentation Outline – NOTE THE DETAILED OUTLINE', level=1)
    
    doc.add_heading('1. INTRODUCTION: THE EPHEMERAL AI FALLACY (5 mins)', level=2)
    doc.add_paragraph("• The Mirage of Deletion: Contrast cloud-side deletion logic with local disk remnants. Explain the LevelDB LSM-tree append-only structure.", style='List Bullet')
    doc.add_paragraph("• Threat Modeling: Assess how malicious actors, insider threats, and infostealer malware harvest unencrypted local browser caches to steal corporate intellectual property.", style='List Bullet')
    doc.add_paragraph("• Target Mapping: Location of IndexedDB LevelDB instances for ChatGPT, Claude, and Gemini on Windows AppData and macOS Library.", style='List Bullet')
    
    doc.add_heading('2. REVERSING CHROMIUM LEVELDB & STORAGE SCHEMA STRUCTURES (5 mins)', level=2)
    doc.add_paragraph("• ChatGPT (Web & Desktop): Layout of keys and IndexedDB schemas storing conversation content, model attributes, and temporary states.", style='List Bullet')
    doc.add_paragraph("• Claude & Gemini: Structural differences in LevelDB records, and mapping active workspace telemetry.", style='List Bullet')
    doc.add_paragraph("• Session Credentials: Identifying unencrypted session tokens (e.g. JWTs and session IDs) stored in IndexedDB that allow session hijacking.", style='List Bullet')
    
    doc.add_heading('3. ZERO-DEPENDENCY CARVING OF LEVELDB & COMPRESSED SSTABLES (10 mins)', level=2)
    doc.add_paragraph("• Production Constraints: Why native C++ leveldb drivers cannot be compiled/installed on endpoint machines during triage.", style='List Bullet')
    doc.add_paragraph("• Pure-Python Parsers: Building a pure-Python Protocol Buffer varint32 reader and block boundary parser.", style='List Bullet')
    doc.add_paragraph("• Snappy Decompression: Rebuilding a Snappy decompressor (literals and copy tags) in Python to extract compressed data blocks directly from SSTables.", style='List Bullet')
    doc.add_paragraph("• Active Lock Bypass: Opening write-ahead logs (.log) in read-only binary mode to clone records without tripping OS locks.", style='List Bullet')
    
    doc.add_heading('4. REBUILDING CONVERSATION TREES & DECODING V8 SERIALIZATION (10 mins)', level=2)
    doc.add_paragraph("• V8 Deserialization format: Decoding JavaScript serialized values.", style='List Bullet')
    doc.add_paragraph("• String Anomalies: Solving the OneByteString (ASCII) vs TwoByteString (UTF-16LE) tag length calculation byte bug.", style='List Bullet')
    doc.add_paragraph("• Nesting & Boundaries: Handling padding bytes and tracking nesting depth (object/array tags) to prevent premature parsing termination on 0x7b.", style='List Bullet')
    doc.add_paragraph("• Role Math: Mapping user vs assistant roles via Smi-shifted index keys and index parity:", style='List Bullet')
    doc.add_paragraph("   role = \"user\" if (index / 2) mod 2 = 1 else \"assistant\"")
    doc.add_paragraph("• Timeline Assembly: Ordering chats using file modification times (mtime) and database byte offsets.", style='List Bullet')
    
    doc.add_heading('5. INCIDENT SCENARIO: THE INSIDER INTELLECTUAL PROPERTY EXFILTRATION (5 mins)', level=2)
    doc.add_paragraph("• Scenario setup: Developer pastes proprietary source code and API keys, deletes the chat, and clears browser history.", style='List Bullet')
    doc.add_paragraph("• Forensic Acquisition: Bypassing locks, cloning raw log files, and executing the Verity carver.", style='List Bullet')
    doc.add_paragraph("• Reconstruction: Demonstrating how the entire code block and thread context are recovered with zero-data loss.", style='List Bullet')
    
    doc.add_heading('6. INCIDENT SCENARIO: SESSION HIJACKING & FORENSIC INTEGRITY (5 mins)', level=2)
    doc.add_paragraph("• The Session Takeover: Extracting unencrypted session tokens from IndexedDB caches and hijacking the user's active session without entering credentials.", style='List Bullet')
    doc.add_paragraph("• Chain of Custody: Creating a tamper-proof forensic envelope by signing canonical JSON payloads using HMAC-SHA256.", style='List Bullet')
    doc.add_paragraph("• Integrity Verification: Validating signature seals in the dashboard using the Web Cryptography API.", style='List Bullet')
    
    doc.add_heading('7. CONCLUSION & MITIGATIONS (3 mins)', level=2)
    doc.add_paragraph("• Key takeaways: V8 deserialization, Snappy decompression, and browser multi-profile path audits.", style='List Bullet')
    doc.add_paragraph("• Mitigations: Implementing active cache eviction policies on logout, disk encryption (BitLocker), and EDR monitoring rules.", style='List Bullet')
    doc.add_paragraph("• Open Source Release: Announcing the public GitHub release of the Verity framework.", style='List Bullet')
    
    doc.add_heading('Questions & Answers', level=1)
    doc.add_heading('Is This Content New (Not Previously Presented/Published)?', level=2)
    doc.add_paragraph("Yes. While general SQLite/LevelDB structures are documented, the reverse engineering of client-side IndexedDB databases for LLM interfaces and the creation of a zero-dependency, pure-Python Snappy block carver for deleted data reconstruction is brand-new research.")
    
    doc.add_heading('Have You/Do You Plan to Submit This Talk to Another Conference?', level=2)
    doc.add_paragraph("No.")
    
    doc.add_heading('What new research, concept, technique, or approach is included in your submission?', level=2)
    doc.add_paragraph("1. Low-Level V8 Forensics of GenAI Clients: We document the client-side IndexedDB database schemas of ChatGPT, Claude, and Gemini, proving that cleared conversations remain on disk.", style='List Bullet')
    doc.add_paragraph("2. Zero-Dependency Snappy/SSTable Carver: We introduce a technique to parse LevelDB SSTables and decompress Snappy blocks entirely in memory using pure Python, eliminating compiled C++ dependencies.", style='List Bullet')
    doc.add_paragraph("3. Dynamic Profile Sweeping & Lock Bypass: Automatically identifies all active Chromium browser profiles and reads write-ahead logs in read-only mode to bypass OS file locks.", style='List Bullet')
    
    doc.add_heading('Provide 3 Audience Takeaways.', level=2)
    doc.add_paragraph("1. AI Local Footprint Map: Deep technical knowledge of where and how local AI client data is stored, and the specific database keys containing session credentials and chat logs.", style='List Bullet')
    doc.add_paragraph("2. LevelDB V8 Carving Mechanics: The ability to write compile-free scripts to parse Protocol Buffer varints, Snappy blocks, and V8 serialized strings from raw disk dumps.", style='List Bullet')
    doc.add_paragraph("3. Open-Source Tooling: Access to the open-source Verity tool to immediately audit endpoints and secure corporate AI workloads.", style='List Bullet')
    
    doc.add_heading('If applicable, what problem does your research solve?', level=2)
    doc.add_paragraph("Traditional EDR agents and forensic suites do not parse GenAI IndexedDB records. If sensitive intellectual property or active AI account credentials are leaked, incident responders lack playbooks to audit the breach. Verity solves this by providing a single-file, zero-dependency Python script that bypasses locks, extracts browser profiles, and compiles verified reports under a cryptographic chain of custody.")
    
    doc.add_heading('Will You Be Releasing a New Tool? If Yes, Describe the Tool.', level=2)
    doc.add_paragraph("Yes. We will be releasing Verity, an open-source, zero-dependency Python framework designed for incident responders. It automates the extraction and parsing of LLM databases (browser IndexedDB and desktop wrappers), handles live SQLite lock bypasses, performs process-to-socket audits, and outputs a clean forensic evidence report to JSON or a centralized web-based lab interface.")
    
    doc.add_heading('Is This a New Vulnerability? If Yes, Describe the Vulnerability.', level=2)
    doc.add_paragraph("Yes. While not a direct remote code execution exploit in Chromium, our research exposes two distinct local vulnerabilities that directly impact modern enterprise AI usage:")
    doc.add_paragraph("1. Unprotected Local Session Credential Leakage in IndexedDB (Chrome/Edge): Chromium-based browsers fail to apply OS-level cryptographic protections (like DPAPI or Keychain) to cookies and authorization tokens stored within IndexedDB LevelDB instances. Standard cookies and passwords are encrypted, but IndexedDB stores active LLM tokens in plaintext, enabling low-privilege processes or infostealers to hijack AI sessions.", style='List Bullet')
    doc.add_paragraph("2. Forensic Leakage of Deleted Telemetry (LevelDB LSM-Tree Persistence): LevelDB's append-only design means that clicking 'Delete Chat' in an LLM web UI does not purge the data from disk. It merely deletes the database index pointer. The actual chat messages, source code blocks, and sensitive credentials remain on-disk in write-ahead logs (.log) and uncompacted Sorted String Tables (.sst/.ldb) indefinitely. We show how these can be forensically reconstructed using raw byte carving.", style='List Bullet')
    
    doc.add_heading('Will Your Presentation Include a Demo? If Yes, Describe the Demo.', level=2)
    doc.add_paragraph("Yes. The presentation will include a pre-recorded demo video showing:")
    doc.add_paragraph("1. A developer pasting proprietary corporate source code into ChatGPT under Chrome Profile 3, and then deleting the conversation using the web interface.", style='List Bullet')
    doc.add_paragraph("2. The execution of the Verity script on the target workstation.", style='List Bullet')
    doc.add_paragraph("3. The dynamic discovery of Profile 3 databases and the instant forensic reassembly of the 'deleted' code blocks by carving the LevelDB data files.", style='List Bullet')
    doc.add_paragraph("4. The extraction of active session tokens to hijack the developer's session, followed by the verification of the evidence integrity using the Web Cryptography API.", style='List Bullet')
    
    doc.add_heading('Provide the Names of the Speakers Presenting and Their Previous Speaking Experience.', level=2)
    doc.add_paragraph("Speaker: Dr. Sapna Vikram Mewundi (Lead Forensics Researcher, Verity Project)\nPrevious Speaking Experience: Regular presenter at local security meetups, DEF CON Groups, and contributor to open-source forensic tools. First-time speaker at Black Hat Briefings.")
    
    doc.add_heading('Does Your Company/ Employer Provide a Solution to the Issue Addressed? If Yes, Please Provide Details.', level=2)
    doc.add_paragraph("No.")
    
    doc.add_heading('Do You Want to Provide a White Paper for Review by the Board?', level=2)
    doc.add_paragraph("Yes. We will provide our comprehensive whitepaper detailing the exact binary layouts, Snappy decompression code, and timeline reassembly algorithms.")
    
    doc.save(docx_path)
    print("DOCX build completed successfully.")

if __name__ == '__main__':
    build_docx()
